#testing how to take data from word table as list


import pandas as pd
from docx import Document
from tkinter import filedialog , ttk
import tkinter as tk



def create_list_from_word_table(number_of_table):
 
 global df
 iteration_of_each_table = document.tables[number_of_table]
 data = [[cell.text for cell in row.cells] for row in iteration_of_each_table.rows]
 df = pd.DataFrame(data)
 dataframe_list = df.values.tolist()
 dataframe_list.remove(dataframe_list[0])

 def verify_integer_list(lista):
       if all( isinstance(item , int) for item in lista) or all( isinstance(item , float) for item in lista) or all( isinstance(item , complex) for item in lista) :
         return True
       else:
         return False

 for i in range(0 , len(dataframe_list)):
   if verify_integer_list(dataframe_list[i]):
      dataframe_list.remove(dataframe_list[i])

 return dataframe_list

 
 

 
def create_list_from_all_word_tables():

 all_dataframe_tables_list = []
 for number_of_the_table in range(number_of_tables):
   iteration_of_each_table = document.tables[number_of_the_table]
   data = [[cell.text for cell in row.cells] for row in iteration_of_each_table.rows]
   df = pd.DataFrame(data)
   dataframe_list = df.values.tolist()
   dataframe_list.remove(dataframe_list[0])

   def verify_integer_list(lista):
       if all( isinstance(item , int) for item in lista) or all( isinstance(item , float) for item in lista) or all( isinstance(item , complex) for item in lista) :
         return True
       else:
         return False

   for i in range(0 , len(dataframe_list)):
    if verify_integer_list(dataframe_list[i]):
      dataframe_list.remove(dataframe_list[i])
   

   all_dataframe_tables_list.append(dataframe_list)
   
 return all_dataframe_tables_list

 


def import_word_file():

  global document , number_of_tables
  word_file = filedialog.askopenfilename(initialdir = "/" , title = "Select a File" , filetypes = (("Word Documents" , "*.docx*" ) , ("all files","*.*") ))
  document = Document(word_file)
  number_of_tables = len(document.tables)


  new_window = tk.Tk()
  new_window.geometry("280x220")
  new_window.title("Choose Word Table")
  new_window.iconbitmap("chart.ico")
  new_window.config(bg = "#9fb49f")
  new_window.resizable(0,0)

  table_positions = number_of_tables
  cmb_tables_var = tk.StringVar()
  cmb_tables = ttk.Combobox(new_window , textvariable=cmb_tables_var , state= "readonly")

  cmb_tables['values'] = tuple( ["Choose table number"] + ["All tables"] + list(cmb_tables['values']) )
  for x in range(table_positions):
      cmb_tables['values'] = tuple(list(cmb_tables['values']) + [int(x+1)])
  cmb_tables.place(relx = 0.5 , rely = 0.5 , anchor=tk.CENTER)

  def cmb_function(event):
    if cmb_tables_var.get() != "Choose table number" and cmb_tables_var.get() != "All tables" :
            create_list_from_word_table(int(cmb_tables.get())-1)
            new_window.destroy()

    elif cmb_tables_var.get() == "All tables" :
            create_list_from_all_word_tables()
            new_window.destroy()
            

  cmb_tables.current(0)
  cmb_tables.bind("<<ComboboxSelected>>" , cmb_function)
  new_window.mainloop()